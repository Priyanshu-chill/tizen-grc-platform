import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def compile_docx(md_path, docx_path):
    print(f"Reading markdown from: {md_path}")
    if not os.path.exists(md_path):
        print("Error: Markdown file not found.")
        return
        
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    doc = Document()
    
    # Configure margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    in_code_block = False
    in_list = False
    
    print("Writing Word document blocks...")
    for line in lines:
        stripped = line.strip()
        
        # Code block toggle
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            continue
            
        # Headings
        if stripped.startswith("# "):
            title_text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(title_text)
            run.font.name = 'Arial'
            run.font.size = Pt(22)
            run.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            continue
            
        if stripped.startswith("## "):
            h_text = stripped[3:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(h_text)
            run.font.name = 'Arial'
            run.font.size = Pt(16)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            continue
            
        if stripped.startswith("### "):
            h_text = stripped[4:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(h_text)
            run.font.name = 'Arial'
            run.font.size = Pt(13)
            run.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            continue
            
        # Lists
        if stripped.startswith("* ") or stripped.startswith("- "):
            list_text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(list_text)
            continue
            
        # Table rows parsing (simple conversion to text block for formatting simplicity)
        if stripped.startswith("|"):
            if "---" in stripped:
                continue
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            run = p.add_run(" | ".join(cells))
            run.font.size = Pt(10)
            run.italic = True
            continue
            
        if stripped == "":
            continue
            
        # Standard Paragraph
        # Replace inline bold formatting
        clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(clean_text)
        
    doc.save(docx_path)
    print(f"Word document saved to: {docx_path}")


def compile_html(md_path, html_path):
    print(f"Reading markdown from: {md_path}")
    if not os.path.exists(md_path):
        return
        
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()
        
    # Convert markdown to basic HTML elements
    html_body = content
    
    # Headings
    html_body = re.sub(r'^# (.*?)$', r'<h1>\1</h1>', html_body, flags=re.MULTILINE)
    html_body = re.sub(r'^## (.*?)$', r'<h2>\1</h2>', html_body, flags=re.MULTILINE)
    html_body = re.sub(r'^### (.*?)$', r'<h3>\1</h3>', html_body, flags=re.MULTILINE)
    
    # Code Blocks
    html_body = re.sub(r'```(.*?)```', r'<pre><code>\1</code></pre>', html_body, flags=re.DOTALL)
    
    # Tables (simple conversion)
    lines = html_body.split('\n')
    in_table = False
    table_lines = []
    new_lines = []
    
    for line in lines:
        if line.strip().startswith('|'):
            if '---' in line:
                continue
            in_table = True
            cells = [f"<td>{c.strip()}</td>" for c in line.split('|')[1:-1]]
            table_lines.append("<tr>" + "".join(cells) + "</tr>")
        else:
            if in_table:
                new_lines.append("<table class='table-doc'>" + "".join(table_lines) + "</table>")
                table_lines = []
                in_table = False
            new_lines.append(line)
            
    html_body = "\n".join(new_lines)
    
    # Lists
    html_body = re.sub(r'^\* (.*?)$', r'<li>\1</li>', html_body, flags=re.MULTILINE)
    html_body = re.sub(r'^- (.*?)$', r'<li>\1</li>', html_body, flags=re.MULTILINE)
    
    # Paragraph elements around raw text (excluding html blocks)
    # Inline Bold
    html_body = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', html_body)
    
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>UTSGRCP v3.0 Thesis Documentation</title>
    <style>
        body {{
            font-family: 'Inter', -apple-system, sans-serif;
            color: #334155;
            line-height: 1.6;
            background-color: #f8fafc;
            margin: 0;
            padding: 40px;
        }}
        .container {{
            max-width: 900px;
            margin: 0 auto;
            background: #ffffff;
            padding: 50px;
            border-radius: 12px;
            box-shadow: 0 4px 6px -1px rgb(0 0 0 / 0.1);
        }}
        h1 {{
            color: #0f172a;
            font-size: 32px;
            border-bottom: 2px solid #e2e8f0;
            padding-bottom: 12px;
            margin-top: 30px;
        }}
        h2 {{
            color: #1e293b;
            font-size: 22px;
            margin-top: 30px;
            border-left: 4px solid #3b82f6;
            padding-left: 12px;
        }}
        h3 {{
            color: #334155;
            font-size: 16px;
            margin-top: 20px;
        }}
        pre {{
            background: #f1f5f9;
            padding: 15px;
            border-radius: 6px;
            overflow-x: auto;
            border: 1px solid #e2e8f0;
        }}
        code {{
            font-family: 'Courier New', Courier, monospace;
            font-size: 13px;
            color: #0f172a;
        }}
        .table-doc {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}
        .table-doc td {{
            padding: 10px;
            border: 1px solid #e2e8f0;
            font-size: 14px;
        }}
        .table-doc tr:nth-child(even) {{
            background: #f8fafc;
        }}
        .print-btn {{
            background: #3b82f6;
            color: #ffffff;
            border: none;
            padding: 10px 20px;
            border-radius: 6px;
            font-weight: 600;
            cursor: pointer;
            margin-bottom: 30px;
        }}
        .print-btn:hover {{
            background: #2563eb;
        }}
        @media print {{
            body {{ padding: 0; background: #fff; }}
            .container {{ box-shadow: none; padding: 0; }}
            .print-btn {{ display: none; }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <button class="print-btn" onclick="window.print()">Print / Save as PDF</button>
        {html_body}
    </div>
</body>
</html>
"""
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"HTML printable document saved to: {html_path}")


if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    md_file = os.path.join(base_dir, "README_THESIS.md")
    docx_file = os.path.join(base_dir, "README_THESIS.docx")
    html_file = os.path.join(base_dir, "README_THESIS.html")
    
    compile_docx(md_file, docx_file)
    compile_html(md_file, html_file)
